"""Patient data persistence and management."""

import json
import os
import re
from datetime import datetime, date
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple


SUPPORTED_EXTENSIONS = {".pdf", ".jpg", ".jpeg", ".png", ".docx", ".json", ".txt", ".xml"}

# ─── Document category age limits ────────────────────────────────────────────

_CATEGORY_ENV = {
    "analyzes":    ("DOC_MAX_DAYS_ANALYSES",   30),
    "researches":  ("DOC_MAX_DAYS_RESEARCHES", 730),
    "inspections": ("DOC_MAX_DAYS_INSPECTIONS", 180),
}

_RU_MONTHS = {
    "январ": 1, "феврал": 2, "март": 3, "апрел": 4,
    "май": 5, "мая": 5, "июн": 6, "июл": 7, "август": 8,
    "сентябр": 9, "октябр": 10, "ноябр": 11, "декабр": 12,
}


def extract_document_date(text: str) -> Optional[datetime]:
    """Extract the most prominent date from document text content.

    Tries several common formats used in Russian medical documents:
      - DD.MM.YYYY  / DD/MM/YYYY
      - YYYY-MM-DD  (ISO)
      - DD <russian-month> YYYY  (e.g. "01 января 2026")
    Searches the first 3 000 characters first (header area), then the full text.
    Returns the first valid date found, or None.
    """
    # Limit to reasonable portion for speed; scan header first
    search_areas = [text[:3000], text] if len(text) > 3000 else [text]

    patterns = [
        # DD.MM.YYYY or DD/MM/YYYY
        (r"\b(\d{1,2})[./](\d{1,2})[./](20\d{2})\b",
         lambda m: _safe_date(int(m.group(3)), int(m.group(2)), int(m.group(1)))),
        # YYYY-MM-DD
        (r"\b(20\d{2})-(\d{2})-(\d{2})\b",
         lambda m: _safe_date(int(m.group(1)), int(m.group(2)), int(m.group(3)))),
        # DD <ru-month> YYYY  (e.g. "01 января 2026" or "1 март 2025")
        (r"\b(\d{1,2})\s+([а-яёА-ЯЁ]{3,10})\s+(20\d{2})\b",
         lambda m: _parse_ru_date(m.group(1), m.group(2), m.group(3))),
    ]

    for area in search_areas:
        for pattern, converter in patterns:
            for m in re.finditer(pattern, area, re.IGNORECASE):
                result = converter(m)
                if result is not None:
                    return result
    return None


def _safe_date(year: int, month: int, day: int) -> Optional[datetime]:
    try:
        return datetime(year, month, day)
    except ValueError:
        return None


def _parse_ru_date(day_str: str, month_str: str, year_str: str) -> Optional[datetime]:
    month_lower = month_str.lower()
    month_num = None
    for prefix, num in _RU_MONTHS.items():
        if month_lower.startswith(prefix):
            month_num = num
            break
    if month_num is None:
        return None
    return _safe_date(int(year_str), month_num, int(day_str))


# Prefixes that indicate the analysis_text contains an API/transport error
# rather than a real document summary.  Comparison is done case-insensitively
# on the *stripped* beginning of the text.
_ERROR_ANALYSIS_PREFIXES = (
    "**rate limit",
    "**api error",
    "**error (",
    "**error:",
)


def is_error_analysis(analysis_text: Optional[str]) -> bool:
    """Return True when analysis_text holds an error message instead of a summary.

    Detects known API error patterns (rate limits, HTTP errors, etc.).
    Does NOT flag legitimate "content not extractable" messages for binary files,
    since re-analysing those would produce the same result anyway.
    """
    if not analysis_text or not analysis_text.strip():
        return True
    lower = analysis_text.strip().lower()
    return any(lower.startswith(prefix) for prefix in _ERROR_ANALYSIS_PREFIXES)


def _category_max_days(doc_path: Path) -> Optional[int]:
    """Return the configured max-age in days for a document, or None if no limit."""
    # Find the top-level subfolder under documents/
    parts = doc_path.parts
    try:
        doc_idx = next(i for i, p in enumerate(parts) if p == "documents")
    except StopIteration:
        return None
    if doc_idx + 1 >= len(parts):
        return None
    category = parts[doc_idx + 1]
    if category not in _CATEGORY_ENV:
        return None
    env_var, default = _CATEGORY_ENV[category]
    raw = os.getenv(env_var, "")
    try:
        return int(raw)
    except (ValueError, TypeError):
        return default


class PatientDataHandler:
    """Manages patient data storage and retrieval."""

    def __init__(self, project_root: str = "."):
        self.project_root = Path(project_root)
        self.patient_data_dir = self.project_root / "patient-data"
        self.sessions_dir = self.patient_data_dir / "sessions"
        self.documents_dir = self.project_root / "documents"
        self.current_patient_file = self.patient_data_dir / "current-patient.md"

        self.patient_data_dir.mkdir(parents=True, exist_ok=True)
        self.sessions_dir.mkdir(parents=True, exist_ok=True)

    # ─── Patient record ────────────────────────────────────────────────────

    def create_or_load_patient(self, patient_id: Optional[str] = None) -> Dict[str, Any]:
        """Create a blank patient record."""
        if patient_id is None:
            patient_id = f"Patient_{datetime.now().strftime('%Y_%m_%d_%H%M%S')}"
        return {
            "patient_id": patient_id,
            "diagnostic_start_date": datetime.now().strftime("%d.%m.%Y"),
            "last_updated": datetime.now().strftime("%d.%m.%Y %H:%M"),
            "chief_complaints": [],
            "associated_symptoms_present": [],
            "associated_symptoms_absent": [],
            "history_of_present_illness": "",
            "chronic_conditions": "",
            "surgeries_injuries": "",
            "allergy_history": "",
            "epidemiological_history": "",
            "family_history": "",
            "medications": "",
            "habits": "",
            "lifestyle": "",
            "red_flags": [],
            "documents_analyzed": [],
            "diagnosis_list": [],
            "diagnostic_status": "In progress — anamnesis collection",
        }

    def save_patient_data(
        self,
        patient_data: Dict[str, Any],
        structured_anamnesis: Optional[Dict[str, Any]] = None,
    ) -> None:
        """Save patient data to current-patient.md in structured format per anamnesis.md rules."""
        # Merge structured anamnesis into patient_data if provided
        if structured_anamnesis:
            for key, value in structured_anamnesis.items():
                if value:
                    patient_data[key] = value

        content = self._format_patient_md(patient_data)
        with open(self.current_patient_file, "w", encoding="utf-8") as f:
            f.write(content)

    def _format_patient_md(self, data: Dict[str, Any]) -> str:
        """Format patient record as structured Markdown per anamnesis.md rules."""
        lines = []
        pid = data.get("patient_id", "Unknown")
        lines.append(f"# Анамнез пациента [{pid}]\n")
        lines.append(f"**Дата начала:** {data.get('diagnostic_start_date', 'N/A')}")
        lines.append(f"**Последнее обновление:** {datetime.now().strftime('%d.%m.%Y %H:%M')}\n")

        # 1. Chief complaints — table format per anamnesis.md
        lines.append("## 1. Основные жалобы\n")
        complaints = data.get("chief_complaints", [])
        if complaints and isinstance(complaints[0], dict):
            lines.append("| Симптом | Локализация | Интенсивность | Длительность | Начало | Динамика | Провоцирующие факторы | Облегчающие факторы |")
            lines.append("|---------|-------------|---------------|--------------|--------|----------|-----------------------|---------------------|")
            for c in complaints:
                lines.append(
                    f"| {c.get('symptom','')} | {c.get('location','')} | {c.get('intensity','')} "
                    f"| {c.get('duration','')} | {c.get('onset','')} | {c.get('dynamics','')} "
                    f"| {c.get('triggers','')} | {c.get('relieving_factors','')} |"
                )
        elif complaints:
            for c in complaints:
                lines.append(f"- {c}")
        else:
            lines.append("_(не собраны)_")
        lines.append("")

        # 2. Associated symptoms — present and absent (negative)
        lines.append("## 2. Сопутствующие симптомы\n")
        present = data.get("associated_symptoms_present", [])
        absent = data.get("associated_symptoms_absent", [])
        lines.append("**Присутствующие:**")
        if present:
            for s in present:
                lines.append(f"- {s}")
        else:
            lines.append("- _(не собраны)_")
        lines.append("")
        lines.append("**Отсутствующие (негативные симптомы):**")
        if absent:
            for s in absent:
                lines.append(f"- {s}")
        else:
            lines.append("- _(не собраны)_")
        lines.append("")

        # 3. History of present illness
        lines.append("## 3. Анамнез настоящего заболевания\n")
        lines.append(data.get("history_of_present_illness", "_(не собран)_"))
        lines.append("")

        # 4. Past medical history — table
        lines.append("## 4. Анамнез жизни\n")
        lines.append("| Параметр | Данные |")
        lines.append("|----------|--------|")
        pmh_fields = [
            ("Хронические заболевания", "chronic_conditions"),
            ("Операции и травмы", "surgeries_injuries"),
            ("Аллергоанамнез", "allergy_history"),
            ("Эпиданамнез", "epidemiological_history"),
            ("Семейный анамнез", "family_history"),
            ("Принимаемые препараты", "medications"),
            ("Вредные привычки", "habits"),
            ("Образ жизни", "lifestyle"),
        ]
        for label, key in pmh_fields:
            value = data.get(key) or "—"
            lines.append(f"| {label} | {value} |")
        lines.append("")

        # 5. Red flags
        red_flags = data.get("red_flags", [])
        if red_flags:
            lines.append("## 5. Красные флаги ⚠️\n")
            for flag in red_flags:
                lines.append(f"- ⚠️ **{flag}**")
            lines.append("")

        # 6. Documents analyzed
        docs = data.get("documents_analyzed", [])
        if docs:
            lines.append("## 6. Анализ медицинских документов\n")
            for doc in docs:
                lines.append(f"### {doc.get('filename', 'Документ')}")
                lines.append(f"**Тип:** {doc.get('doc_type', '—')} | **Дата документа:** {doc.get('doc_date', '—')}\n")
                analysis_text = doc.get("analysis_text", "")
                if analysis_text:
                    lines.append(analysis_text)
                lines.append("")

        # 7. Diagnosis list
        dx_list = data.get("diagnosis_list", [])
        if dx_list:
            lines.append("## 7. Список диагнозов (вероятности)\n")
            lines.append("| Диагноз | Вероятность | Уверенность | Обоснование |")
            lines.append("|---------|-------------|-------------|-------------|")
            for dx in dx_list:
                lines.append(
                    f"| {dx.get('name','')} | {dx.get('probability','')} "
                    f"| {dx.get('confidence','')} | {dx.get('rationale','')} |"
                )
            lines.append("")

        lines.append(f"**Статус:** {data.get('diagnostic_status', 'В процессе')}")
        return "\n".join(lines)

    # ─── Single session file ───────────────────────────────────────────────

    def save_session(
        self,
        patient_id: str,
        stage_number: int,
        messages: List[Dict[str, str]],
        patient_data: Dict[str, Any],
        structured_anamnesis: Optional[Dict[str, Any]] = None,
        documents_analyzed: Optional[List[Dict[str, Any]]] = None,
        extra: Optional[Dict[str, Any]] = None,
    ) -> None:
        """Save complete session to a single JSON file (overwrites each time)."""
        session_file = self.sessions_dir / f"session_{patient_id}.json"

        # Preserve original start time if file already exists
        started_at = datetime.now().isoformat()
        if session_file.exists():
            try:
                with open(session_file, "r", encoding="utf-8") as f:
                    existing = json.load(f)
                started_at = existing.get("started_at", started_at)
            except Exception:
                pass

        record = {
            "patient_id": patient_id,
            "started_at": started_at,
            "updated_at": datetime.now().isoformat(),
            "stage_number": stage_number,
            "messages": messages,
            "patient_data": patient_data,
            "structured_anamnesis": structured_anamnesis or {},
            "documents_analyzed": documents_analyzed or [],
            **(extra or {}),
        }
        with open(session_file, "w", encoding="utf-8") as f:
            json.dump(record, f, indent=2, ensure_ascii=False)

    def load_latest_session(self) -> Optional[Dict[str, Any]]:
        """Load the most recently modified session file."""
        # New format: session_{patient_id}.json
        session_files = list(self.sessions_dir.glob("session_*.json"))
        if session_files:
            latest = max(session_files, key=lambda f: f.stat().st_mtime)
            with open(latest, "r", encoding="utf-8") as f:
                data = json.load(f)
            # Normalize key for backward compat (old files used "saved_at")
            if "saved_at" in data and "updated_at" not in data:
                data["updated_at"] = data["saved_at"]
            return data

        return None

    # ─── Document scanning ─────────────────────────────────────────────────

    def scan_documents(self) -> List[Path]:
        """Recursively find all supported medical documents in documents/ folder.

        Applies age filtering per document category (see scan_documents_filtered).
        Returns only accepted documents.
        """
        accepted, _ = self.scan_documents_filtered()
        return accepted

    def scan_documents_filtered(self) -> Tuple[List[Path], List[Dict[str, Any]]]:
        """Return (accepted, skipped) after applying per-category age filters.

        Age limits are configured via env vars:
          DOC_MAX_DAYS_ANALYSES   (default 30)   — documents/analyzes/
          DOC_MAX_DAYS_RESEARCHES (default 730)  — documents/researches/
          DOC_MAX_DAYS_INSPECTIONS (default 180) — documents/inspections/

        The document date is extracted from its text content, NOT from file mtime.
        If no date can be extracted, the document is accepted (safe default).

        Each skipped entry contains:
          path, doc_date (str DD.MM.YYYY), age_days, max_days, category
        """
        if not self.documents_dir.exists():
            return [], []

        all_files = sorted(
            p for p in self.documents_dir.rglob("*")
            if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
        )

        today = datetime.now().date()
        accepted: List[Path] = []
        skipped: List[Dict[str, Any]] = []

        for path in all_files:
            max_days = _category_max_days(path)
            if max_days is None:
                accepted.append(path)
                continue

            content = self.read_document_content(path)
            doc_dt = extract_document_date(content) if content else None

            if doc_dt is None:
                # Can't determine date — accept by default
                accepted.append(path)
                continue

            age_days = (today - doc_dt.date()).days
            if age_days <= max_days:
                accepted.append(path)
            else:
                # Determine human-readable category label
                parts = path.parts
                try:
                    doc_idx = next(i for i, p in enumerate(parts) if p == "documents")
                    cat = parts[doc_idx + 1] if doc_idx + 1 < len(parts) else "?"
                except StopIteration:
                    cat = "?"
                skipped.append({
                    "path": path,
                    "filename": path.name,
                    "doc_date": doc_dt.strftime("%d.%m.%Y"),
                    "age_days": age_days,
                    "max_days": max_days,
                    "category": cat,
                })

        return accepted, skipped

    # ─── Text / JSON ───────────────────────────────────────────────────────

    def read_document_content(self, file_path: Path) -> Optional[str]:
        """Read text content from a document file.

        Supported formats and extraction methods:
          .txt / .xml / .json  — direct UTF-8 read
          .pdf (text layer)    — pypdf
          .pdf (scanned)       — pymupdf rendering + Tesseract OCR
          .jpg / .jpeg / .png  — Tesseract OCR via pytesseract
          .docx                — python-docx paragraph extraction

        Returns None only when extraction truly fails (missing library AND no fallback).
        When OCR is used the result is prefixed with [OCR] so the LLM is aware
        that recognition quality may vary.
        """
        suffix = file_path.suffix.lower()

        if suffix in (".txt", ".xml"):
            try:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    return f.read(8000)
            except Exception:
                return None

        if suffix == ".json":
            try:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    return f.read(8000)
            except Exception:
                return None

        if suffix == ".pdf":
            return self._read_pdf(file_path)

        if suffix in (".jpg", ".jpeg", ".png"):
            return self._ocr_image(file_path)

        if suffix == ".docx":
            return self._read_docx(file_path)

        return None

    # ─── PDF ───────────────────────────────────────────────────────────────

    def _read_pdf(self, file_path: Path) -> Optional[str]:
        """Try pypdf for text-layer PDFs; fall back to OCR for scanned pages."""
        text = self._pypdf_text(file_path)
        if text:
            return text
        # Scanned PDF: render pages with pymupdf and OCR
        return self._ocr_pdf_pages(file_path)

    def _pypdf_text(self, file_path: Path) -> Optional[str]:
        """Extract text from a PDF using pypdf. Returns None if text layer is absent."""
        try:
            import pypdf  # type: ignore
            reader = pypdf.PdfReader(str(file_path))
            pages_text = []
            for i, page in enumerate(reader.pages[:15]):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    pages_text.append(f"[Страница {i+1}]\n{page_text}")
            combined = "\n\n".join(pages_text)
            # Require at least 100 chars to consider the text layer meaningful
            return combined[:8000] if len(combined.strip()) >= 100 else None
        except ImportError:
            return None
        except Exception:
            return None

    def _ocr_pdf_pages(self, file_path: Path) -> Optional[str]:
        """Render PDF pages with PyMuPDF and OCR each page with Tesseract."""
        try:
            import pymupdf  # type: ignore  (package: pymupdf)
        except ImportError:
            try:
                import fitz as pymupdf  # type: ignore  (older package name)
            except ImportError:
                return None  # PyMuPDF not installed

        try:
            doc = pymupdf.open(str(file_path))
        except Exception:
            return None

        pages_text = []
        for i, page in enumerate(doc[:15]):
            # Render at 200 DPI for decent OCR quality
            mat = pymupdf.Matrix(200 / 72, 200 / 72)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_bytes = pix.tobytes("png")
            page_text = self._tesseract_bytes(img_bytes)
            if page_text and page_text.strip():
                pages_text.append(f"[Страница {i+1}]\n{page_text}")

        doc.close()
        if not pages_text:
            return None
        combined = "[OCR]\n\n" + "\n\n".join(pages_text)
        return combined[:8000]

    # ─── Images ────────────────────────────────────────────────────────────

    def _ocr_image(self, file_path: Path) -> Optional[str]:
        """Run Tesseract OCR on an image file (JPG/PNG)."""
        try:
            from PIL import Image  # type: ignore
        except ImportError:
            return None  # Pillow not installed

        try:
            img = Image.open(str(file_path))
            text = self._tesseract_pil(img)
            if not text or not text.strip():
                return None
            return ("[OCR]\n\n" + text)[:8000]
        except Exception:
            return None

    # ─── DOCX ──────────────────────────────────────────────────────────────

    def _read_docx(self, file_path: Path) -> Optional[str]:
        """Extract text from a DOCX file using python-docx."""
        try:
            from docx import Document  # type: ignore
        except ImportError:
            return None  # python-docx not installed

        try:
            doc = Document(str(file_path))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        parts.append(" | ".join(row_cells))
            combined = "\n".join(parts)
            return combined[:8000] if combined.strip() else None
        except Exception:
            return None

    # ─── Tesseract helpers ─────────────────────────────────────────────────

    @staticmethod
    def _tesseract_pil(img: "Image") -> Optional[str]:  # type: ignore
        """Run pytesseract on a PIL Image. Tries rus+eng, falls back to eng."""
        try:
            import pytesseract  # type: ignore
        except ImportError:
            return None

        for lang in ("rus+eng", "eng"):
            try:
                return pytesseract.image_to_string(img, lang=lang)
            except Exception:
                continue
        return None

    @staticmethod
    def _tesseract_bytes(img_bytes: bytes) -> Optional[str]:
        """Run pytesseract on raw PNG bytes (from PyMuPDF pixmap)."""
        try:
            from PIL import Image  # type: ignore
            import io
            img = Image.open(io.BytesIO(img_bytes))
            return PatientDataHandler._tesseract_pil(img)
        except Exception:
            return None

    def save_document_analysis(
        self,
        patient_id: str,
        documents: List[Dict[str, Any]],
    ) -> None:
        """Append document analysis results to current-patient.md and session."""
        # Update the session file with doc analysis
        session_file = self.sessions_dir / f"session_{patient_id}.json"
        if session_file.exists():
            try:
                with open(session_file, "r", encoding="utf-8") as f:
                    record = json.load(f)
                record["documents_analyzed"] = documents
                record["updated_at"] = datetime.now().isoformat()
                with open(session_file, "w", encoding="utf-8") as f:
                    json.dump(record, f, indent=2, ensure_ascii=False)
            except Exception:
                pass

    # ─── Case archiving ────────────────────────────────────────────────────

    def archive_patient(self, patient_id: str) -> Path:
        """Move completed case to patient-data/archive/ per patient-data-management.md rules.

        Copies current-patient.md and session file to archive, then resets
        current-patient.md so the system is ready for a new patient.
        Returns the path of the archived file.
        """
        archive_dir = self.patient_data_dir / "archive"
        archive_dir.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        archive_name = f"{patient_id}_{timestamp}"

        # Archive current-patient.md
        archived_md = archive_dir / f"{archive_name}.md"
        if self.current_patient_file.exists():
            archived_md.write_text(
                self.current_patient_file.read_text(encoding="utf-8"),
                encoding="utf-8",
            )

        # Archive session JSON
        session_file = self.sessions_dir / f"session_{patient_id}.json"
        if session_file.exists():
            archived_json = archive_dir / f"{archive_name}_session.json"
            archived_json.write_text(
                session_file.read_text(encoding="utf-8"),
                encoding="utf-8",
            )

        self.update_diagnostic_log(patient_id, f"Case archived → {archived_md.name}")
        return archived_md

    # ─── Document analysis cache ──────────────────────────────────────────

    @property
    def _cache_file(self) -> Path:
        return self.patient_data_dir / "document-cache.json"

    def _load_cache(self) -> Dict[str, Any]:
        """Load the full cache dict (keyed by filename)."""
        if not self._cache_file.exists():
            return {}
        try:
            with open(self._cache_file, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}

    def _save_cache(self, cache: Dict[str, Any]) -> None:
        with open(self._cache_file, "w", encoding="utf-8") as f:
            json.dump(cache, f, indent=2, ensure_ascii=False)

    def get_cached_document(
        self, filename: str, file_path: Optional[Path] = None
    ) -> Optional[Dict[str, Any]]:
        """Return cached analysis for filename, or None if not found / file changed.

        If file_path is given, compares current mtime against cached mtime —
        returns None (cache miss) when the file has been modified since last analysis.
        """
        cache = self._load_cache()
        entry = cache.get(filename)
        if entry is None:
            return None
        if file_path and file_path.exists():
            cached_mtime = entry.get("file_mtime")
            current_mtime = file_path.stat().st_mtime
            if cached_mtime is not None and abs(current_mtime - float(cached_mtime)) > 1:
                return None  # file changed → re-analyze
        return entry

    def save_to_document_cache(
        self, filename: str, result: Dict[str, Any], file_path: Optional[Path] = None
    ) -> None:
        """Persist a document analysis result to the cache."""
        cache = self._load_cache()
        entry = dict(result)
        entry["analyzed_at"] = datetime.now().isoformat()
        if file_path and file_path.exists():
            entry["file_mtime"] = file_path.stat().st_mtime
        cache[filename] = entry
        self._save_cache(cache)

    # ─── Diagnostic log ────────────────────────────────────────────────────

    def update_diagnostic_log(self, patient_id: str, entry: str) -> None:
        """Append an entry to memory/diagnostic-log.md."""
        log_file = self.project_root / "memory" / "diagnostic-log.md"
        log_file.parent.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        log_entry = f"\n- **{timestamp}** | {patient_id}: {entry}\n"
        if log_file.exists():
            with open(log_file, "a", encoding="utf-8") as f:
                f.write(log_entry)
        else:
            with open(log_file, "w", encoding="utf-8") as f:
                f.write(f"# Diagnostic Log\n{log_entry}")
